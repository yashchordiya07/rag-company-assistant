"""
rag/document_loader.py
─────────────────────────────────────────────────────────────
Document ingestion layer for the RAG pipeline.

Supports loading:
    • PDF   — via pypdf
    • DOCX  — via python-docx
    • TXT   — plain UTF-8 / Latin-1
    • MD    — plain text (Markdown)

Each loaded document is returned as a ``Document`` dataclass that
carries the raw text, per-page metadata, and a SHA-256 fingerprint
for deduplication.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from utils.helpers import clean_text, file_hash, get_file_extension, new_uuid
from utils.logger import get_logger

log = get_logger(__name__)


# ─── Domain model ────────────────────────────────────────────
@dataclass
class PageContent:
    """Content and metadata for a single page / section."""

    text: str
    page_number: int | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class Document:
    """
    Represents a fully loaded source document.

    Attributes
    ----------
    doc_id      : Unique identifier (UUID)
    source_path : Original file path
    file_name   : Base filename
    file_type   : Extension (pdf | docx | txt | md)
    pages       : List of PageContent objects
    full_text   : Concatenated text of all pages (cleaned)
    file_hash   : SHA-256 of the raw file bytes
    metadata    : Arbitrary key/value pairs from the loader
    """

    doc_id: str
    source_path: str
    file_name: str
    file_type: str
    pages: list[PageContent]
    full_text: str
    file_hash: str
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def num_pages(self) -> int:
        return len(self.pages)

    @property
    def char_count(self) -> int:
        return len(self.full_text)


# ─── Loader base ─────────────────────────────────────────────
class BaseLoader:
    """Abstract document loader interface."""

    def load(self, path: Path, data: bytes | None = None) -> Document:
        raise NotImplementedError


# ─── PDF Loader ──────────────────────────────────────────────
class PDFLoader(BaseLoader):
    """Load PDF files using pypdf."""

    def load(self, path: Path, data: bytes | None = None) -> Document:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ImportError("Install pypdf: pip install pypdf") from exc

        raw = data or path.read_bytes()
        reader = PdfReader(io.BytesIO(raw))

        pages: list[PageContent] = []
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = clean_text(text)
            if text.strip():
                pages.append(
                    PageContent(
                        text=text,
                        page_number=page_num,
                        metadata={
                            "page": page_num,
                            "total_pages": len(reader.pages),
                        },
                    )
                )

        full_text = "\n\n".join(p.text for p in pages)
        doc_metadata = {}
        if reader.metadata:
            doc_metadata = {
                "title": reader.metadata.get("/Title", ""),
                "author": reader.metadata.get("/Author", ""),
                "creator": reader.metadata.get("/Creator", ""),
            }

        log.info(f"PDF loaded: {path.name} — {len(pages)} pages")
        return Document(
            doc_id=new_uuid(),
            source_path=str(path),
            file_name=path.name,
            file_type="pdf",
            pages=pages,
            full_text=full_text,
            file_hash=_bytes_hash(raw),
            metadata=doc_metadata,
        )


# ─── DOCX Loader ─────────────────────────────────────────────
class DOCXLoader(BaseLoader):
    """Load DOCX files using python-docx."""

    def load(self, path: Path, data: bytes | None = None) -> Document:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise ImportError("Install python-docx: pip install python-docx") from exc

        raw = data or path.read_bytes()
        docx = DocxDocument(io.BytesIO(raw))

        # Group paragraphs into logical "pages" of ~50 paragraphs each
        PARA_PER_PAGE = 50
        all_paras = [p.text for p in docx.paragraphs if p.text.strip()]
        pages: list[PageContent] = []

        for i in range(0, len(all_paras), PARA_PER_PAGE):
            chunk_paras = all_paras[i : i + PARA_PER_PAGE]
            text = clean_text("\n".join(chunk_paras))
            page_num = (i // PARA_PER_PAGE) + 1
            if text.strip():
                pages.append(
                    PageContent(
                        text=text,
                        page_number=page_num,
                        metadata={"section": page_num},
                    )
                )

        # Also extract tables
        table_texts: list[str] = []
        for table in docx.tables:
            rows = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    rows.append(row_text)
            if rows:
                table_texts.append("\n".join(rows))

        if table_texts:
            pages.append(
                PageContent(
                    text=clean_text("\n\n".join(table_texts)),
                    page_number=None,
                    metadata={"content_type": "table"},
                )
            )

        full_text = "\n\n".join(p.text for p in pages)
        # Extract core properties
        cp = docx.core_properties
        doc_metadata = {
            "title": cp.title or "",
            "author": cp.author or "",
            "created": str(cp.created) if cp.created else "",
        }

        log.info(f"DOCX loaded: {path.name} — {len(all_paras)} paragraphs")
        return Document(
            doc_id=new_uuid(),
            source_path=str(path),
            file_name=path.name,
            file_type="docx",
            pages=pages,
            full_text=full_text,
            file_hash=_bytes_hash(raw),
            metadata=doc_metadata,
        )


# ─── Plain-text Loader ───────────────────────────────────────
class TextLoader(BaseLoader):
    """Load TXT and MD files as a single page."""

    def load(self, path: Path, data: bytes | None = None) -> Document:
        raw = data or path.read_bytes()

        # Attempt common encodings gracefully
        text = ""
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                text = raw.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        text = clean_text(text)
        ext = get_file_extension(path)
        pages = [PageContent(text=text, page_number=1, metadata={"encoding": "utf-8"})]

        log.info(f"Text loaded: {path.name} — {len(text)} chars")
        return Document(
            doc_id=new_uuid(),
            source_path=str(path),
            file_name=path.name,
            file_type=ext,
            pages=pages,
            full_text=text,
            file_hash=_bytes_hash(raw),
            metadata={},
        )


# ─── Document Loader (factory + dispatcher) ──────────────────
class DocumentLoader:
    """
    High-level document loader that dispatches to the correct
    loader based on file extension.

    Usage
    -----
    loader = DocumentLoader()
    doc = loader.load("./data/hr_policy.pdf")
    """

    _LOADERS: dict[str, type[BaseLoader]] = {
        "pdf": PDFLoader,
        "docx": DOCXLoader,
        "txt": TextLoader,
        "md": TextLoader,
    }

    def __init__(self) -> None:
        self._instances: dict[str, BaseLoader] = {}

    def _get_loader(self, ext: str) -> BaseLoader:
        if ext not in self._LOADERS:
            raise ValueError(
                f"Unsupported file type: .{ext}. "
                f"Supported: {list(self._LOADERS)}"
            )
        if ext not in self._instances:
            self._instances[ext] = self._LOADERS[ext]()
        return self._instances[ext]

    def load(self, path: str | Path, data: bytes | None = None) -> Document:
        """
        Load a document from a file path or raw bytes.

        Parameters
        ----------
        path : str | Path
            File path (used for metadata and extension detection).
        data : bytes | None
            If provided, raw file bytes (skips reading from disk).

        Returns
        -------
        Document
        """
        path = Path(path)
        ext = get_file_extension(path)
        loader = self._get_loader(ext)

        try:
            doc = loader.load(path, data)
            log.debug(
                f"Loaded document",
                file=path.name,
                type=ext,
                chars=doc.char_count,
                pages=doc.num_pages,
            )
            return doc
        except Exception as exc:
            log.error(f"Failed to load {path.name}: {exc}")
            raise

    def load_directory(
        self, directory: str | Path, recursive: bool = False
    ) -> list[Document]:
        """
        Load all supported documents from a directory.

        Parameters
        ----------
        directory : str | Path
            Root directory to scan.
        recursive : bool
            If True, scan subdirectories as well.
        """
        directory = Path(directory)
        if not directory.is_dir():
            raise NotADirectoryError(f"{directory} is not a directory")

        pattern = "**/*" if recursive else "*"
        paths = [
            p
            for p in directory.glob(pattern)
            if p.is_file() and get_file_extension(p) in self._LOADERS
        ]

        log.info(f"Found {len(paths)} supported files in {directory}")
        docs: list[Document] = []
        for path in paths:
            try:
                docs.append(self.load(path))
            except Exception as exc:
                log.warning(f"Skipping {path.name}: {exc}")

        return docs


# ─── Private helpers ─────────────────────────────────────────
def _bytes_hash(data: bytes) -> str:
    import hashlib
    return hashlib.sha256(data).hexdigest()
